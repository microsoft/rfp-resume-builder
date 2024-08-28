
import os
from dotenv import load_dotenv


from langchain_openai import AzureChatOpenAI

from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient, BlobClient
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult, AnalyzeDocumentRequest

from typing import List, Dict

from io import BytesIO
from pypandoc.pandoc_download import download_pandoc
import json
import pypandoc

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from helper_functions import get_rfp_analysis_from_db
from prompts import reorder_work_experience_prompt
from docx2pdf import convert
import tempfile
import os
import contextlib
import tempfile
import time





load_dotenv()

# Azure Blob Storage
connect_str = os.getenv("STORAGE_ACCOUNT_CONNECTION_STRING")
storage_account_name = os.getenv("STORAGE_ACCOUNT_NAME")
container_name = "resumes"

#Doc intelligence
form_recognizer_endpoint = os.getenv("FORM_RECOGNIZER_ENDPOINT")
form_recognizer_key = os.getenv("FORM_RECOGNIZER_KEY")



# Azure OpenAI
aoai_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
aoai_key = os.getenv("AZURE_OPENAI_API_KEY")
aoai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")

endpoint = form_recognizer_endpoint
credential = AzureKeyCredential(form_recognizer_key)
document_intelligence_client = DocumentIntelligenceClient(endpoint, credential)



primary_llm = AzureChatOpenAI(
    azure_deployment=aoai_deployment,
    api_version="2024-05-01-preview",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    api_key=aoai_key,
    azure_endpoint=aoai_endpoint
)

primary_llm_json = AzureChatOpenAI(
    azure_deployment=aoai_deployment,
    api_version="2024-05-01-preview",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    api_key=aoai_key,
    azure_endpoint=aoai_endpoint,
    model_kwargs={"response_format": {"type": "json_object"}}
)

work_experience_prompt = """You are an AI assistant. You are given a resume in PDF format. Your job is to read the resume and output the structured information in JSON format. 
It is critical that you don't add/change/remove anything. You must only extract the information that is already present in the resume verbatim in all cases.

#Formatting Guidance#

Output valid json with the following fields: 

1. 'analysis' - this is the only output you will make that isn't in the resume content. It should be a brief analysis of the resume structure. Take note on where you see work experience and "Prior to CDM" experience.
1. 'name': The name of the candidate.
2. 'title': The title/role of the candidate.
3. 'work_experience': A list of projects. Needs to capture only the projects worked on at CDM Smith, nothing under the "Prior to CDM Smith" section.

#Example#
{
  "analysis": "The resume has a clear structure with a work experience section that lists projects worked on at CDM Smith and a section for prior to CDM Smith.",
  "name": "John Doe",
  "title": "Senior Engineer",
  "work_experience": [
    {
      "title": "NEW PROJECT TITLE 1",
      "description": "NEW PROJECT DESCRIPTION 1"
    },
    {
      "title": "NEW PROJECT TITLE 2",
      "description": "NEW PROJECT DESCRIPTION 2"
    }
  ]
}



"""

insertion_prompt = """
    Analyze the following resume text and identify where the work experience section begins.
    Return a JSON object with two keys:
    1. 'analysis': Comment on where you think the right substring to split is located at. We will be inserting BEFORE the split substring you identify, so consider that. 
    2. 'start_phrase': The substring to split on

    We will be inserting a new project BEFORE this substring. So think about what substring you would need to identify in order to insert the new project in the right place. 

    ###Examples###

    User: Mr. Abhay Ashok is a licensed architect who has experience working in Government projects like Treasury, Hospitals, Masterplans and other commercial sectors like Museum Designs, dental clinics etc. His software skills include Revit, AutoCAD, Sketchup, Photoshop, Adobe Premier Pro, Lumion, Enscape.  
Work Experience 
Architect, Fort Meyer Beach, USA 2023. Mr. Abhay was responsible in developing the phase wise 4D sequencing for the Wastewater Treatment Project using Lumion and Premier Pro. The company has won the pursuit and Abhay and the team received Encore – Option 9 Award recognized by McKim Tanner for delivering visuals for the Pursuit. 

    Assistant: {
        "analysis": "I can clearly see the work experience section heading. I need to identify a substring that we can insert directly above to add the new project in the right place. If i output 'Work Experience',
          the new project will be inserted above that which is not correct. If i output 'Architect, Fort Meyer Beach, USA 2023', the new project will be inserted above that which is would produce a logical resume with the new project underneath work experience.",
        "start_phrase": "Architect, Fort Meyer Beach, USA 2023"
    }

    """


def find_insert_position(doc):
    # Extract text from the document
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Prompt for the LLM
   
    
    messages = [
        {"role": "system", "content": insertion_prompt},
        {"role": "user", "content": full_text}
    ]
    
    result = primary_llm_json.invoke(messages)
    result_json = json.loads(result.content)
    
    print("Analysis:", result_json['analysis'])
    print("Start Phrase:", result_json['start_phrase'])

    return result_json['start_phrase']

def insert_new_project(doc, new_project, insert_phrase):
    for para in doc.paragraphs:
        if insert_phrase in para.text:
            # Insert the new project before the paragraph containing the insert phrase
            new_para = para.insert_paragraph_before()
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = new_para.add_run(new_project['title'])
            run.bold = True
            run.font.size = Pt(11)
            new_para.add_run('\n' + new_project['description'])
            new_para.style = 'Normal'

            # Add a blank line after the new project
            para.insert_paragraph_before()
            

            print(f"New project inserted before: '{para.text}'")
            return True
    
    print(f"Insert phrase '{insert_phrase}' not found in the document.")
    return False

@contextlib.contextmanager
def temporary_file(suffix=None):
    """Context manager for creating and cleaning up a temporary file."""
    fd, path = tempfile.mkstemp(suffix=suffix)
    try:
        os.close(fd)
        yield path
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass

def enhance_resume(resume_name, rfp_name):
    print(f"Enhancing resume {resume_name} for RFP {rfp_name}")
    
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    container_client = blob_service_client.get_container_client(container_name)
    
    folder = "processed/"
    blob_name = folder + resume_name
    blob_client = container_client.get_blob_client(blob_name)

    if not blob_client.exists():
        print(f"Blob {blob_name} not found in {container_name}.")
        return None, None

    print(f"Blob {blob_name} found in {container_name}.")
    
    # Download the blob content
    blob_data = blob_client.download_blob().readall()
    
    # Load the document
    doc = Document(BytesIO(blob_data))

    # Find insert position using LLM
    insert_phrase = find_insert_position(doc)
    print(f"Inserting new project before: '{insert_phrase}'")

    # New projects to insert
    new_projects = [
        {
            "title": "Extremely Relevant Project Title 1",
            "description": "Worked on an extremely important project that is highly relevant to the RFP"
        },
        {
            "title": "Extremely Relevant Project Title 2",
            "description": "Worked on an extremely important project that is highly relevant to the RFP"
        }
    ]

    # Insert the new projects
    for project in new_projects:
        insert_new_project(doc, project, insert_phrase)

    # Prepare file names
    resume_name_without_ext = os.path.splitext(resume_name)[0]
    enhanced_name = f"{resume_name_without_ext}_enhanced"
    enhanced_docx_name = f"{enhanced_name}.docx"
    enhanced_pdf_name = f"{enhanced_name}.pdf"

    try:
        with temporary_file(suffix='.docx') as temp_docx_path, temporary_file(suffix='.pdf') as temp_pdf_path:
            # Save the enhanced DOCX to a temporary file
            doc.save(temp_docx_path)
            print(f"Enhanced resume (DOCX) saved to temporary file: {temp_docx_path}")

            # Add a 1-second delay
            time.sleep(1)
            enhanced_folder = "enhanced/"
            # Convert DOCX to PDF
            try:
                convert(temp_docx_path, temp_pdf_path)
                print(f"Enhanced resume (PDF) converted to temporary file: {temp_pdf_path}")
                # Upload the PDF file to blob storage
                enhanced_pdf_blob_name = enhanced_folder + enhanced_pdf_name
                enhanced_pdf_client = container_client.get_blob_client(enhanced_pdf_blob_name)
                with open(temp_pdf_path, "rb") as pdf_file:
                    enhanced_pdf_client.upload_blob(pdf_file, overwrite=True)
                print(f"Enhanced resume (PDF) uploaded as {enhanced_pdf_blob_name}")
            except Exception as e:
                print(f"Failed to convert DOCX to PDF: {str(e)}")
    

            # Upload the DOCX file to blob storage
            
            enhanced_docx_blob_name = enhanced_folder + enhanced_docx_name
            enhanced_docx_client = container_client.get_blob_client(enhanced_docx_blob_name)
            with open(temp_docx_path, "rb") as docx_file:
                enhanced_docx_client.upload_blob(docx_file, overwrite=True)
            print(f"Enhanced resume (DOCX) uploaded as {enhanced_docx_blob_name}")

            

        return enhanced_docx_blob_name

    except Exception as e:
        print(f"Failed to enhance resume: {str(e)}")
        return None, None


#Main function
if __name__ == "__main__":

    resume_name = "Abhay Ashok 717569.docx"
    #resume_name = "a rajasekar 702292"
   # resume_name = "Acker George 717841"
    #resume_name = "Adam Frank 87487"
    # resume_name = "adams timothy 709049"
    rfp_name = "Combined_Synopsis_and_Solicitation.pdf"
    
    enhance_resume(resume_name, rfp_name)

    
    

